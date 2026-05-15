import os
import gc
import easyocr
import fitz
from pdf2image import convert_from_path, pdfinfo_from_path
from docx import Document
from docx.shared import Pt as DocxPt, Inches
from utils import get_poppler_path, get_model_path
from PIL import Image

def process_ocr_to_word(input_file, output_path, status_callback, stop_event, use_gpu, use_ocr, dpi):
    is_pdf = input_file.lower().endswith('.pdf')
    doc = Document()
    
    # === 智慧極速模式 (非 OCR) ===
    if not use_ocr and is_pdf:
        status_callback("⚡ 啟動極速文字提取模式...", 0.1)
        fitz_doc = fitz.open(input_file)
        total_pages = len(fitz_doc)
        
        for i in range(total_pages):
            if stop_event.is_set(): break
            status_callback(f"⚡ 正在極速寫入 Word (第 {i+1} / {total_pages} 頁)...", 0.1 + 0.8 * ((i+1)/total_pages))
            
            page = fitz_doc[i]
            # 插入圖片對照
            pix = page.get_pixmap(dpi=dpi)
            temp_img = f"temp_fast_word_{i}.jpg"
            pix.save(temp_img)
            
            doc.add_paragraph(f"--- 第 {i+1} 頁 原始圖片 ---").style.font.name = 'Microsoft JhengHei'
            doc.add_picture(temp_img, width=Inches(6.0))
            os.remove(temp_img)
            
            # 直接抓取原生文字
            text = page.get_text("text").strip()
            doc.add_paragraph(f"--- 第 {i+1} 頁 提取文字 ---").style.font.name = 'Microsoft JhengHei'
            if text:
                p = doc.add_paragraph(text)
                p.style.font.name = 'Microsoft JhengHei'
                for run in p.runs: run.font.size = DocxPt(11)
            
            if i < total_pages - 1: doc.add_page_break()
        fitz_doc.close()
        
    # === 傳統 OCR 圖片辨識模式 ===
    else:
        status_callback("⏳ 正在初始化 OCR 模型...", 0.05)
        reader = easyocr.Reader(['ch_tra', 'en'], model_storage_directory=get_model_path(), download_enabled=False, gpu=use_gpu)
        poppler = get_poppler_path()
        
        if is_pdf: total_pages = pdfinfo_from_path(input_file, poppler_path=poppler)["Pages"]
        else: img_cache = [Image.open(input_file).convert('RGB')]; total_pages = 1

        for i in range(total_pages):
            if stop_event.is_set(): break
            status_callback(f"🔍 正在辨識 Word 圖文 (第 {i+1} / {total_pages} 頁)...", 0.1 + 0.8 * ((i+1)/total_pages))
            
            if is_pdf: page_img = convert_from_path(input_file, dpi=dpi, first_page=i+1, last_page=i+1, poppler_path=poppler)[0]
            else: page_img = img_cache[0]
                
            temp_img = f"temp_word_page_{i}.jpg"
            page_img.save(temp_img, 'JPEG')
            
            doc.add_paragraph(f"--- 第 {i+1} 頁 原始圖片 ---").style.font.name = 'Microsoft JhengHei'
            doc.add_picture(temp_img, width=Inches(6.0))
            
            result = reader.readtext(temp_img, paragraph=True)
            doc.add_paragraph(f"--- 第 {i+1} 頁 辨識文字 ---").style.font.name = 'Microsoft JhengHei'
            for bbox, text in result:
                p = doc.add_paragraph(text)
                p.style.font.name = 'Microsoft JhengHei'
                for run in p.runs: run.font.size = DocxPt(11)
            
            if i < total_pages - 1: doc.add_page_break() 
            os.remove(temp_img)
            if is_pdf: del page_img; gc.collect()
    
    if not stop_event.is_set():
        status_callback("💾 正在寫入 Word 檔案...", 0.95)
        doc.save(output_path)
